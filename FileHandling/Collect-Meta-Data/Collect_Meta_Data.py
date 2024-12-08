import os
import pytesseract
import pdfplumber
import pandas as pd
from PIL import Image
from docx import Document
from datetime import datetime
import logging
import xlrd
import cv2
import numpy as np
import regex as re

# Set up logging for debugging purposes
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')

def get_xl_properties(filepath):
    """Extracts properties from Excel files."""
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(0)
        properties = {
            'Sheets in Workbook': len(workbook.sheets()),
            'Rows in first sheet': sheet.nrows,
            'Columns in first sheet': sheet.ncols,
            'Excel File Created': datetime(*xlrd.xldate_as_tuple(workbook.datemode, 0)).strftime('%Y-%m-%d %H:%M:%S')
        }
        return properties
    except Exception as e:
        logging.error(f"Error reading Excel file {filepath}: {e}")
        return {}

def extract_file_metadata(filepath):
    """Extracts file and dynamic metadata."""
    metadata = {}

    try:
        # General metadata
        stats = os.stat(filepath)
        metadata['Size'] = f"{stats.st_size} bytes"
        metadata['Date modified'] = datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
        metadata['Date created'] = datetime.fromtimestamp(stats.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
        metadata['Date accessed'] = datetime.fromtimestamp(stats.st_atime).strftime('%Y-%m-%d %H:%M:%S')
        metadata['File Path'] = os.path.abspath(filepath)
        metadata['File extension'] = os.path.splitext(filepath)[-1]
        metadata['Filename'] = os.path.basename(filepath)
        metadata['Folder Path'] = os.path.dirname(filepath)

        # Dynamic metadata
        if filepath.endswith(".txt"):
            with open(filepath, 'r') as file:
                content = file.read()
                metadata['Word Count'] = len(content.split())
        elif filepath.endswith(".pdf"):
            with pdfplumber.open(filepath) as pdf:
                metadata['Page Count'] = len(pdf.pages)
        elif filepath.endswith(".docx"):
            doc = Document(filepath)
            metadata['Paragraph Count'] = len(doc.paragraphs)
        elif filepath.endswith(('.xls', '.xlsx')):
            metadata.update(get_xl_properties(filepath))
        elif filepath.endswith(".zip"):
            metadata['Zip Archive'] = 'Contains compressed files'
        else:
            metadata['File Type'] = 'Unknown'

    except Exception as e:
        logging.error(f"Error extracting metadata from {filepath}: {e}")
    
    return metadata

def summarize_file_content(filepath):
    """Generate a one-line summary of the file content based on its type."""
    summary = ""
    try:
        if filepath.endswith(".txt"):
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read().replace("\n", " ").strip()  # Replace line breaks with spaces
                summary = f"Text File: {content}" if content else "Empty Text File"
        elif filepath.endswith(".pdf"):
            with pdfplumber.open(filepath) as pdf:
                first_page = pdf.pages[0].extract_text()
                summary = f"PDF Document: {first_page}" if first_page else "Empty PDF File"
        elif filepath.endswith(".docx"):
            doc = Document(filepath)
            first_paragraph = doc.paragraphs[0].text if doc.paragraphs else ""
            summary = f"Word Document: {first_paragraph}" if first_paragraph else "Empty Word Document"
        elif filepath.endswith(".xls"):
            summary = "Excel File: Contains spreadsheet data"
        elif filepath.endswith(".zip"):
            summary = "ZIP Archive: Contains compressed files"
        else:
            summary = f"File of type {os.path.splitext(filepath)[-1]}: No content preview available"
    except Exception as e:
        logging.error(f"Error summarizing content of {filepath}: {e}")
        summary = "Error generating summary"
    return summary

def extract_text_from_image(file_path):
    """Extracts text from image files using pytesseract (OCR)."""
    try:
        image = Image.open(file_path)

        # Convert to numpy array for OpenCV processing
        img = np.array(image)
        
        # Apply preprocessing for better OCR results
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        threshold = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                          cv2.THRESH_BINARY, 11, 2)
        dilated = cv2.dilate(threshold, None, iterations=1)

        # Perform OCR
        text = pytesseract.image_to_string(dilated)

        # Ensure to return cleaned text
        return text.strip() if text else "No text found in image"
    except Exception as e:
        logging.error(f"Error extracting text from image {file_path}: {e}")
        return f"Error: {str(e)}"

def dynamic_classification(content):
    """Dynamically classify content based on keywords or structures."""
    labels = []
    
    # Define common keywords and patterns
    keywords_dict = {
    "invoice": r"(invoice|bill|receipt|order|total|due)",
    "contract": r"(contract|agreement|terms|conditions|signed)",
    "resume": r"(resume|cv|curriculum vitae|experience|education)",
    "report": r"(report|analysis|summary|findings|conclusion)",
    "bank_statement": r"(bank statement|account statement|transaction history|deposit|withdrawal)",
    "tax_return": r"(tax return|tax form|tax document|income tax|sales tax)",
    "insurance_policy": r"(insurance policy|policy document|coverage|claim|premium)",
    "investment_statement": r"(investment statement|portfolio statement|stock|bond|mutual fund)",
    "mortgage_statement": r"(mortgage statement|loan statement|payment|interest rate|amortization)",
    "lease_agreement": r"(lease agreement|rental agreement|lease term|rent|security deposit)",
    "power_of_attorney": r"(power of attorney|POA|authorization|proxy|delegate)",
    "will": r"(will|testament|inheritance|bequest|estate)",
    "deed": r"(deed|property deed|land deed|title deed|conveyance)",
    "court_order": r"(court order|judge's order|judicial order|ruling|decree)",
    "medical_record": r"(medical record|patient record|health record|medical history|chart)",
    "prescription": r"(prescription|prescription drug|medication|pharmacy|doctor's order)",
    "lab_report": r"(lab report|test result|medical test|blood test|diagnostic test)",
    "hospital_bill": r"(hospital bill|medical bill|hospital charges|medical expenses|healthcare costs)",
    "insurance_claim": r"(insurance claim|claim form|medical claim|accident claim|health claim)",
    "research_paper": r"(research paper|academic paper|scholarly article|thesis|dissertation)",
    "syllabus": r"(syllabus|course outline|curriculum|course schedule|class plan)",
    "transcript": r"(transcript|academic transcript|grade report|GPA|academic record)",
    "certificate": r"(certificate|diploma|degree|license|qualification)",
    "letter_of_recommendation": r"(letter of recommendation|recommendation letter|reference letter|endorsement)",
    "business_plan": r"(business plan|business proposal|business strategy|market plan|financial plan)",
    "marketing_plan": r"(marketing plan|marketing strategy|marketing campaign|advertising|promotion)",
    "project_plan": r"(project plan|project proposal|project schedule|project timeline|project scope)",
    "IT_ticket": r"(IT ticket|help desk ticket|support ticket|incident report|problem report)",
    "performance_review": r"(performance review|performance appraisal|employee review|performance evaluation)",
    "passport": r"(passport|travel document|visa|citizenship|nationality)",
    "driver's_license": r"(driver'sws license|driving license|driver's permit|DL)",
    "birth_certificate": r"(birth certificate|birth record|certificate of live birth)",
    "marriage_certificate": r"(marriage certificate|marriage license)",
    "death_certificate": r"(death certificate|death record)",
    "email": r"(email|electronic mail|message|inbox|outbox)",
    "letter": r"(letter|formal letter|informal letter|business letter|personal letter)",
    "memo": r"(memo|memorandum|note|reminder|announcement)",
    "report": r"(report|analysis|summary|findings|conclusion)",
    "presentation": r"(presentation|slide deck|PowerPoint|Keynote|slide show)",
    "budget": r"(budget|expense report|invoice|receipt|payment|check|credit card statement)",
    "legal": r"(contract|agreement|lease|deed|will|power of attorney|court order|subpoena)",
    "medical": r"(prescription|medical record|lab report|x-ray|MRI|CT scan|hospital bill|insurance claim)",
    "academic": r"(syllabus|textbook|research paper|thesis|dissertation|transcript|certificate|diploma)",
    "professional": r"(resume|CV|cover letter|job application|business plan|marketing plan|project plan|IT ticket|performance review)",
    "personal": r"(passport|driver's license|birth certificate|marriage certificate|death certificate|tax return|insurance policy)",
    "general": r"(email|letter|memo|report|presentation|document|file|folder|PDF|Word|Excel|PowerPoint)"
}
    
    # Check content against each pattern
    for label, pattern in keywords_dict.items():
        if re.search(pattern, content, re.IGNORECASE):
            labels.append(label)

    if not labels:
        labels.append("Document")
    
    return labels

def extract_text_from_file(file_path):
    """Extracts text based on file type (txt, pdf, docx, csv, image, etc.)."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)

    elif ext == '.docx':
        return extract_text_from_docx(file_path)

    elif ext in ['.csv', '.xls', '.xlsx']:
        return extract_text_from_spreadsheet(file_path)

    elif ext in ['.jpg', '.jpeg', '.png']:
        return extract_text_from_image(file_path)

    else:
        raise ValueError(f"Unsupported file type: {ext}")

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file using pdfplumber."""
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        return str(e)

def extract_text_from_docx(file_path):
    """Extracts text from a Word document (DOCX)."""
    try:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return str(e)

def extract_text_from_spreadsheet(file_path):
    """Extracts text from CSV, XLS, or XLSX files using pandas."""
    try:
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
        elif file_path.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file_path)
            return df.to_string(index=False)
    except Exception as e:
        return str(e)

def print_metadata(metadata):
    """Print out non-empty metadata."""
    for key, value in metadata.items():
        if value:  # Only print if value exists
            print(f"\t{key} = {value}")

if __name__ == '__main__':
    # Get the file path from the user input
    file_path = input("Please enter the file path: ")

    # Check if the file exists
    if os.path.exists(file_path):
        print(f"Processing file: {file_path}")

        # Extract and print file metadata
        file_metadata = extract_file_metadata(file_path)
        print_metadata(file_metadata)

        # Extract text content from file
        file_content = extract_text_from_file(file_path)

        # Dynamically classify the file based on its content
        file_labels = dynamic_classification(file_content)
        print(f"\nDynamic Classification Labels: {', '.join(file_labels)}")

        # Generate and print a summary of the file content
        file_summary = summarize_file_content(file_path)
        print(f"\nSummary: {file_summary}")
    else:
        print("The file does not exist.")
