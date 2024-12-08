import os
import pytesseract
import pdfplumber
import pandas as pd
from PIL import Image
from docx import Document
from datetime import datetime
import logging
import xlrd
import cv2
import numpy as np
import regex as re

# Set up logging for debugging purposes
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')

def get_xl_properties(filepath):
    """Extracts properties from Excel files."""
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(0)
        properties = {
            'Sheets in Workbook': len(workbook.sheets()),
            'Rows in first sheet': sheet.nrows,
            'Columns in first sheet': sheet.ncols,
            'Excel File Created': datetime(*xlrd.xldate_as_tuple(workbook.datemode, 0)).strftime('%Y-%m-%d %H:%M:%S')
        }
        return properties
    except Exception as e:
        logging.error(f"Error reading Excel file {filepath}: {e}")
        return {}

def extract_file_metadata(filepath):
    """Extracts file and dynamic metadata."""
    metadata = {}

    try:
        # General metadata
        stats = os.stat(filepath)
        metadata['Size'] = f"{stats.st_size} bytes"
        metadata['Date modified'] = datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
        metadata['Date created'] = datetime.fromtimestamp(stats.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
        metadata['Date accessed'] = datetime.fromtimestamp(stats.st_atime).strftime('%Y-%m-%d %H:%M:%S')
        metadata['File Path'] = os.path.abspath(filepath)
        metadata['File extension'] = os.path.splitext(filepath)[-1]
        metadata['Filename'] = os.path.basename(filepath)
        metadata['Folder Path'] = os.path.dirname(filepath)

        # Dynamic metadata
        if filepath.endswith(".txt"):
            with open(filepath, 'r') as file:
                content = file.read()
                metadata['Word Count'] = len(content.split())
        elif filepath.endswith(".pdf"):
            with pdfplumber.open(filepath) as pdf:
                metadata['Page Count'] = len(pdf.pages)
        elif filepath.endswith(".docx"):
            doc = Document(filepath)
            metadata['Paragraph Count'] = len(doc.paragraphs)
        elif filepath.endswith(('.xls', '.xlsx')):
            metadata.update(get_xl_properties(filepath))
        elif filepath.endswith(".zip"):
            metadata['Zip Archive'] = 'Contains compressed files'
        else:
            metadata['File Type'] = 'Unknown'

    except Exception as e:
        logging.error(f"Error extracting metadata from {filepath}: {e}")
    
    return metadata

def dynamic_classification(content):
    """Dynamically classify content based on keywords or structures."""
    labels = []
    keywords_dict = {
        "invoice": r"(invoice|bill|receipt|order|total|due)",
        "contract": r"(contract|agreement|terms|conditions|signed)",
        "resume": r"(resume|cv|curriculum vitae|experience|education)",
        "report": r"(report|analysis|summary|findings|conclusion)",
        "bank_statement": r"(bank statement|account statement|transaction history|deposit|withdrawal)",
        "tax_return": r"(tax return|tax form|tax document|income tax|sales tax)",
        "insurance_policy": r"(insurance policy|policy document|coverage|claim|premium)",
        "investment_statement": r"(investment statement|portfolio statement|stock|bond|mutual fund)",
        "mortgage_statement": r"(mortgage statement|loan statement|payment|interest rate|amortization)",
        "lease_agreement": r"(lease agreement|rental agreement|lease term|rent|security deposit)",
        "power_of_attorney": r"(power of attorney|POA|authorization|proxy|delegate)",
        "will": r"(will|testament|inheritance|bequest|estate)",
        "deed": r"(deed|property deed|land deed|title deed|conveyance)",
        "court_order": r"(court order|judge's order|judicial order|ruling|decree)",
        "medical_record": r"(medical record|patient record|health record|medical history|chart)",
        "prescription": r"(prescription|prescription drug|medication|pharmacy|doctor's order)",
        "lab_report": r"(lab report|test result|medical test|blood test|diagnostic test)",
        "hospital_bill": r"(hospital bill|medical bill|hospital charges|medical expenses|healthcare costs)",
        "insurance_claim": r"(insurance claim|claim form|medical claim|accident claim|health claim)",
        "research_paper": r"(research paper|academic paper|scholarly article|thesis|dissertation)",
        "syllabus": r"(syllabus|course outline|curriculum|course schedule|class plan)",
        "transcript": r"(transcript|academic transcript|grade report|GPA|academic record)",
        "certificate": r"(certificate|diploma|degree|license|qualification)",
        "letter_of_recommendation": r"(letter of recommendation|recommendation letter|reference letter|endorsement)",
        "business_plan": r"(business plan|business proposal|business strategy|market plan|financial plan)",
        "marketing_plan": r"(marketing plan|marketing strategy|marketing campaign|advertising|promotion)",
        "project_plan": r"(project plan|project proposal|project schedule|project timeline|project scope)",
        "IT_ticket": r"(IT ticket|help desk ticket|support ticket|incident report|problem report)",
        "performance_review": r"(performance review|performance appraisal|employee review|performance evaluation)",
        "passport": r"(passport|travel document|visa|citizenship|nationality)",
        "driver's_license": r"(driver's license|driving license|driver's permit|DL)",
        "birth_certificate": r"(birth certificate|birth record|certificate of live birth)",
        "marriage_certificate": r"(marriage certificate|marriage license)",
        "death_certificate": r"(death certificate|death record)",
        "email": r"(email|electronic mail|message|inbox|outbox)",
        "letter": r"(letter|formal letter|informal letter|business letter|personal letter)",
        "memo": r"(memo|memorandum|note|reminder|announcement)",
        "report": r"(report|analysis|summary|findings|conclusion)",
        "presentation": r"(presentation|slide deck|PowerPoint|Keynote|slide show)",
        "budget": r"(budget|expense report|invoice|receipt|payment|check|credit card statement)",
        "legal": r"(contract|agreement|lease|deed|will|power of attorney|court order|subpoena)",
        "medical": r"(prescription|medical record|lab report|x-ray|MRI|CT scan|hospital bill|insurance claim)",
        "academic": r"(syllabus|textbook|research paper|thesis|dissertation|transcript|certificate|diploma)",
        "professional": r"(resume|CV|cover letter|job application|business plan|marketing plan|project plan|IT ticket|performance review)",
        "personal": r"(passport|driver's license|birth certificate|marriage certificate|death certificate|tax return|insurance policy)",
        "general": r"(email|letter|memo|report|presentation|document|file|folder|PDF|Word|Excel|PowerPoint)"
    }

    for label, pattern in keywords_dict.items():
        if re.search(pattern, content, re.IGNORECASE):
            labels.append(label)
    return labels if labels else ["Document"]

def print_metadata(metadata):
    """Print out non-empty metadata."""
    for key, value in metadata.items():
        if value:  # Only print if value exists
            print(f"\t{key} = {value}")

def summarize_txt(file_path):
    """Summarize text files."""
    try:
        with open(file_path, 'r', errors='ignore') as f:
            content = f.read()
            return f"Text file summary: {content[:100]}..."  # Example: First 100 characters
    except Exception as e:
        logging.error(f"Error summarizing text file {file_path}: {e}")
        return "Error summarizing text file."

def summarize_pdf(file_path):
    """Summarize PDF files."""
    try:
        with pdfplumber.open(file_path) as pdf:
            first_page = pdf.pages[0] if pdf.pages else None
            return f"PDF summary: {first_page.extract_text()[:100]}..." if first_page else "PDF is empty."
    except Exception as e:
        logging.error(f"Error summarizing PDF file {file_path}: {e}")
        return "Error summarizing PDF file."

def summarize_docx(file_path):
    """Summarize DOCX files."""
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return f"DOCX summary: {paragraphs[0][:100]}..." if paragraphs else "DOCX file is empty."
    except Exception as e:
        logging.error(f"Error summarizing DOCX file {file_path}: {e}")
        return "Error summarizing DOCX file."

# Main Function
if __name__ == "__main__":
    # Prompt user to input the file path
    file_path = input("Please enter the file path: ")

    if os.path.isfile(file_path):
        # Extract and print metadata
        metadata = extract_file_metadata(file_path)
        print_metadata(metadata)

        # Extract and print content summary
        if file_path.endswith(".txt"):
            summary = summarize_txt(file_path)
        elif file_path.endswith(".pdf"):
            summary = summarize_pdf(file_path)
        elif file_path.endswith(".docx"):
            summary = summarize_docx(file_path)
        else:
            summary = "File type not supported for summary."

        print(f"Summary: {summary}")
    else:
        print(f"The file at {file_path} does not exist.")
