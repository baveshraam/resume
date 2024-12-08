import os
import re
import pdfplumber
import docx
import spacy
import torch
import pandas as pd
import numpy as np
from transformers import BertTokenizer, BertModel
from sklearn.metrics.pairwise import cosine_similarity
import logging
from collections import defaultdict

# Initialize models
nlp = spacy.load("en_core_web_sm")
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
model = BertModel.from_pretrained('bert-base-uncased')

# Set up logging
# Set up logging to print to console
logging.basicConfig(
    level=logging.INFO,  # Ensure INFO level logs are enabled
    format="%(asctime)s - %(message)s",
    handlers=[logging.StreamHandler()]  # Print to console
)


# Utility functions for text extraction
def extract_phone_numbers(text):
    phone_pattern = re.compile(r'\+?[0-9]{1,4}?[-.\s\(\)]?(\(?\d{1,3}?\)?[-.\s]?)?(\d{1,4}[-.\s]?)?\d{1,4}[-.\s]?\d{1,9}')
    return re.findall(phone_pattern, text)

def extract_emails(text):
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b')
    return re.findall(email_pattern, text)

def extract_skills(text):
    skills_keywords = ["Python", "Java", "JavaScript", "SQL", "Machine Learning", "AI", "Data Science", "Git", "Django", "Flask", "C++", "AWS", "Docker"]
    skills_found = [skill for skill in skills_keywords if skill.lower() in text.lower()]
    return skills_found if skills_found else ["N/A"]

def extract_education(text):
    education_patterns = [
        r'(Bachelor|Master|PhD|Doctorate|B\.?S\.?|M\.?S\.?|M\.?B\.?A\.?)\s+in\s+([a-zA-Z\s]+)',
        r'(Degree|Graduation)\s+in\s+([a-zA-Z\s]+)',
        r'([a-zA-Z\s]+)\s+(Degree|Graduate)'
    ]
    education_matches = []
    for pattern in education_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        education_matches.extend(matches)
    return [match[1] for match in education_matches] or ["N/A"]

def extract_work_experience(text):
    work_exp_patterns = [
        r'(?:Work\s*Experience|Professional\s*Experience|Employment\s*History).*?(?:\n\n|\Z)',
        r'([\w\s]+)\s*(?:at|@)\s*([\w\s]+)\s*(?:from|for)\s*(\d{4}(?:-\d{4})?)'
    ]
    experiences = []
    for pattern in work_exp_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
        experiences.extend(matches)
    return experiences[:3] if experiences else ["N/A"]

# Extract basic information from resume text
def extract_comprehensive_resume_info(text):
    name = extract_name(text)
    contact_info = extract_contact_info(text)
    skills = extract_skills(text)
    education = extract_education(text)
    work_experience = extract_work_experience(text)

    return {
        "Name": name,
        "Phone": contact_info["phone_numbers"],
        "Email": contact_info["emails"],
        "Skills": ", ".join(skills),
        "Education": ", ".join(education),
        "Work Experience": str(work_experience),
        "Full Text": text
    }

def extract_name(text):
    doc = nlp(text)
    
    # Initialize variables to track potential names and titles
    possible_names = []
    possible_titles = ["Mr.", "Mrs.", "Ms.", "Dr.", "Prof.", "Sir", "Madam", "Mx."]
    
    # Search for persons and titles in the text
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            possible_names.append(ent.text)
    
    # If no names are detected, try extracting more from sentence patterns
    if not possible_names:
        lines = text.split("\n")
        for line in lines:
            # Look for potential name patterns at the start of a line (like "Dr. John Doe")
            for title in possible_titles:
                if line.startswith(title):
                    possible_names.append(line.strip())
    
    # Handle common cases for name extraction
    if possible_names:
        # Select the longest name (as most likely to be the full name)
        name = max(possible_names, key=len)
        return name
    else:
        # Fallback to a default response if no name found
        return "Name Not Found"

import re

def extract_contact_info(text):
    # Extract phone numbers (with more variations like extensions and international formats)
    phone_pattern = r'(\+?\d{1,3}[\s\.\-]?\(?\d{1,4}\)?[\s\.\-]?\d{1,4}[\s\.\-]?\d{1,4}(?:[\s\.\-]?(ext|x|extension|#)?\s?\d{1,5})?)'
    phone_numbers = re.findall(phone_pattern, text)
    
    # Extract emails (more robust and includes common TLDs)
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    
    # Extract LinkedIn profile URL
    linkedin_pattern = r'https?://(?:www\.)?linkedin\.com/in/([a-zA-Z0-9_-]+)'
    linkedin_profiles = re.findall(linkedin_pattern, text)
    
    # Extract social media handles (Twitter, Instagram, GitHub, Facebook)
    twitter_pattern = r'@([a-zA-Z0-9_]+)'
    instagram_pattern = r'@([a-zA-Z0-9_]+)'  # Instagram handles are similar to Twitter
    github_pattern = r'github\.com/([a-zA-Z0-9_-]+)'
    facebook_pattern = r'https?://(?:www\.)?facebook\.com/([a-zA-Z0-9_.-]+)'
    
    twitter_handles = re.findall(twitter_pattern, text)
    instagram_handles = re.findall(instagram_pattern, text)
    github_handles = re.findall(github_pattern, text)
    facebook_handles = re.findall(facebook_pattern, text)
    
    # Extract physical address (more advanced pattern including zip codes and street names)
    address_pattern = r'\b(\d{1,5}\s(?:[A-Za-z0-9]+(?:[,\.\s])?)+[A-Za-z0-9]{2,},?\s?[A-Za-z]+(?:\s?[A-Za-z]+){0,2},?\s?[A-Za-z]{2,}\s?\d{5}(-\d{4})?)\b'
    addresses = re.findall(address_pattern, text)
    
    # Extract website URLs (for personal websites, portfolios, etc.)
    website_pattern = r'https?://[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    websites = re.findall(website_pattern, text)
    
    # Return all extracted contact information in a structured format
    return {
        "phone_numbers": ", ".join([number[0] for number in phone_numbers]) if phone_numbers else "N/A",
        "emails": ", ".join(emails) if emails else "N/A",
        "linkedin_profile": linkedin_profiles[0] if linkedin_profiles else "N/A",
        "twitter_handles": ", ".join(twitter_handles) if twitter_handles else "N/A",
        "instagram_handles": ", ".join(instagram_handles) if instagram_handles else "N/A",
        "github_handles": ", ".join(github_handles) if github_handles else "N/A",
        "facebook_handles": ", ".join(facebook_handles) if facebook_handles else "N/A",
        "addresses": ", ".join(addresses) if addresses else "N/A",
        "websites": ", ".join(websites) if websites else "N/A"
    }


# PDF, DOCX, and TXT file extraction
import pdfplumber
import pytesseract
from PIL import Image

def process_pdf(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                # Try extracting text
                extracted_text = page.extract_text()
                if extracted_text:  # If text extraction works, append it
                    text += extracted_text
                else:  # If text extraction fails, attempt OCR (Optical Character Recognition)
                    image = page.to_image()
                    ocr_text = pytesseract.image_to_string(image.original)
                    text += ocr_text
            return text
    except Exception as e:
        return f"Error processing PDF: {str(e)}"



import docx

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        # Extract text from headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                text += header.text + "\n"
            for footer in section.footer.paragraphs:
                text += footer.text + "\n"
        
        return text.strip()
    except Exception as e:
        return f"Error processing DOCX: {str(e)}"
def process_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            return text
        except Exception as e:
            return f"Error processing TXT: {str(e)}"
    except Exception as e:
        return f"Error processing TXT: {str(e)}"


# Compute BERT-based similarity between resume and job description
def compute_similarity(text1, text2):
    # Tokenize and encode the texts
    inputs = tokenizer([text1, text2], return_tensors='pt', padding=True, truncation=True, max_length=512)
    
    # Compute embeddings
    with torch.no_grad():
        embeddings = model(**inputs).last_hidden_state.mean(dim=1)
    
    # Compute cosine similarity between the two embeddings
    similarity = cosine_similarity(embeddings[0].numpy().reshape(1, -1), embeddings[1].numpy().reshape(1, -1))
    return similarity[0][0]
def extract_dynamic_keywords(text):
    """
    This function extracts relevant keywords dynamically from the text.
    These could include skills, technologies, roles, etc., that are mentioned in the text.
    """
    # Define a list of potential keywords (can be extended as needed)
    keywords = ["python", "java", "machine learning", "data science", "aws", "cloud", "sql", "git", 
                "docker", "javaScript", "typescript", "c++", "go", "developer", "engineer", "data", 
                "science", "analyst", "systems", "design", "ai", "deep learning", "tensorflow", 
                "pytorch", "flask", "django", "sql", "graphql", "api", "react", "angular", "project manager"]

    # Normalize and extract keywords from text
    found_keywords = [keyword for keyword in keywords if keyword.lower() in text.lower()]
    return found_keywords
def process_resume_matching(job_description, minimum_threshold=0.5):
    folder_path = 'D:/ResumeScreening'
    
    all_candidates = []
    shortlisted_candidates = []
    
    files = [f for f in os.listdir(folder_path) if f.endswith(('.pdf', '.docx', '.txt'))]
    
    for file in files:
        file_path = os.path.join(folder_path, file)
        
        # Log the file being processed
        logging.info(f"Processing file: {file}...")  # Print which file is being processed
        
        if file.endswith(".pdf"):
            resume_text = process_pdf(file_path)
        elif file.endswith(".docx"):
            resume_text = extract_text_from_docx(file_path)
        elif file.endswith(".txt"):
            resume_text = process_txt(file_path)
        
        # Extract name from resume text
        name = extract_name(resume_text)
        
        # Print name of the person whose resume is being processed
        logging.info(f"Processing {name}'s resume...")
        
        resume_info = extract_comprehensive_resume_info(resume_text)
        
        match_result = compare_resume_to_job_description(resume_text, job_description, threshold=minimum_threshold)
        
        resume_info.update({
            "Match Percentage": match_result["match_percentage"] * 100,
            "Skills Match": ", ".join(match_result["skills_match"]),
            "Improvement Suggestions": ", ".join(match_result["improvement_suggestions"]),
            "Suggested Roles": ", ".join(match_result["suggested_roles"]),
            "Similarity Score": match_result["similarity_score"]
        })
        
        all_candidates.append(resume_info)
        
        if match_result["is_match"]:
            shortlisted_candidates.append(resume_info)
    
    all_candidates_df = pd.DataFrame(all_candidates)
    shortlisted_candidates_df = pd.DataFrame(shortlisted_candidates)
    
    all_candidates_df.to_csv('/content/all_candidates.csv', index=False, encoding='utf-8')
    shortlisted_candidates_df.to_csv('/content/shortlisted_candidates.csv', index=False, encoding='utf-8')
    
    return all_candidates_df, shortlisted_candidates_df


def compare_resume_to_job_description(resume_text, job_description, threshold=0.5):
    # Compute similarity score based on BERT embeddings
    similarity_score = compute_similarity(resume_text, job_description)

    # Extract dynamic keywords from the job description
    job_keywords = extract_dynamic_keywords(job_description)

    # Extract keywords from the resume
    resume_keywords = extract_dynamic_keywords(resume_text)

    # Match skills and technologies mentioned in both the resume and job description
    skills_match = list(set(resume_keywords) & set(job_keywords))
    
    # Suggest improvements for missing job description keywords in the resume
    missing_skills = set(job_keywords) - set(resume_keywords)
    improvement_suggestions = [f"Improve knowledge in {skill}" for skill in missing_skills]

    # Role suggestions based on dynamic keywords (if the job description mentions "engineer" or "data", we suggest related roles)
    suggested_roles = []
    if "developer" in job_keywords or "engineer" in job_keywords:
        suggested_roles.append("Software Developer / Engineer")
    if "data" in job_keywords or "science" in job_keywords:
        suggested_roles.append("Data Scientist")
    if "machine learning" in job_keywords or "ai" in job_keywords:
        suggested_roles.append("Machine Learning Engineer")

    # If no roles are suggested, fallback to generic roles
    if not suggested_roles:
        suggested_roles = ["Software Engineer", "Data Scientist", "Machine Learning Engineer"]

    # Compute the result with similarity and matching information
    result = {
        "match_percentage": similarity_score,
        "skills_match": skills_match,
        "improvement_suggestions": improvement_suggestions,
        "suggested_roles": suggested_roles,
        "similarity_score": similarity_score,
        "is_match": similarity_score >= threshold
    }

    return result

# Helper function to print resume details
def print_resume_summary(resume_info):
    logging.info(f"Name: {resume_info['Name']}")
    logging.info(f"Phone: {resume_info['Phone']}")
    logging.info(f"Email: {resume_info['Email']}")
    logging.info(f"Skills: {resume_info['Skills']}")
    logging.info(f"Education: {resume_info['Education']}")
    logging.info(f"Work Experience: {resume_info['Work Experience']}")

# Main entry point for script execution
def main():
    job_description = input("Enter the Job Description: ")
   
    logging.info("Starting resume matching process...")
    all_candidates, shortlisted_candidates = process_resume_matching(job_description, minimum_threshold=0.6)
    
    logging.info(f"Total Candidates: {len(all_candidates)}")
    logging.info(f"Shortlisted Candidates: {len(shortlisted_candidates)}")

if __name__ == "__main__":
    main()
